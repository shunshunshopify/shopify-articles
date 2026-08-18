#!/usr/bin/env python3
from pathlib import Path
from lxml import html
from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_LINE_SPACING, WD_PARAGRAPH_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
SOURCE = ROOT / "drafts" / "ec-site-project-failure-causes.html"
OUTPUT = ROOT / "drafts" / "ec-site-project-failure-causes.docx"

TITLE = "ECサイトプロジェクトはなぜ失敗する？3つの原因と防止策"
META = "ECサイトプロジェクトが失敗する背景を、過大投資、価格競争、公開後の運用不足という3つの原因から解説。利益を残すための予算設計、差別化、運用体制のチェックポイントも紹介します。"
EAST_ASIA_FONT = "Hiragino Sans GB"


def set_east_asia_font(font_element):
    r_pr = font_element._element.get_or_add_rPr()
    r_fonts = r_pr.rFonts
    if r_fonts is None:
        r_fonts = OxmlElement("w:rFonts")
        r_pr.insert(0, r_fonts)
    r_fonts.set(qn("w:eastAsia"), EAST_ASIA_FONT)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for key, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{key}"))
        if node is None:
            node = OxmlElement(f"w:{key}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_borders(table, color="DADCE0", size="4"):
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = borders.find(qn(f"w:{edge}"))
        if tag is None:
            tag = OxmlElement(f"w:{edge}")
            borders.append(tag)
        tag.set(qn("w:val"), "single")
        tag.set(qn("w:sz"), size)
        tag.set(qn("w:space"), "0")
        tag.set(qn("w:color"), color)


def set_table_width(table, widths):
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.first_child_found_in("w:tblW")
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), "9360")
    tbl_w.set(qn("w:type"), "dxa")
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        grid_col = OxmlElement("w:gridCol")
        grid_col.set(qn("w:w"), str(width))
        grid.append(grid_col)
    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            cell.width = Inches(widths[idx] / 1440)
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.first_child_found_in("w:tcW")
            tc_w.set(qn("w:w"), str(widths[idx]))
            tc_w.set(qn("w:type"), "dxa")
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_margins(cell)


def add_hyperlink(paragraph, text, url):
    part = paragraph.part
    rel_id = part.relate_to(url, "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink", is_external=True)
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), rel_id)
    run = OxmlElement("w:r")
    props = OxmlElement("w:rPr")
    fonts = OxmlElement("w:rFonts")
    fonts.set(qn("w:ascii"), "Arial")
    fonts.set(qn("w:hAnsi"), "Arial")
    fonts.set(qn("w:eastAsia"), EAST_ASIA_FONT)
    props.append(fonts)
    color = OxmlElement("w:color")
    color.set(qn("w:val"), "1155CC")
    props.append(color)
    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    props.append(underline)
    run.append(props)
    node = OxmlElement("w:t")
    node.text = text
    run.append(node)
    hyperlink.append(run)
    paragraph._p.append(hyperlink)


def add_inline(paragraph, node):
    if node.text:
        paragraph.add_run(node.text)
    for child in node:
        if child.tag == "a":
            add_hyperlink(paragraph, "".join(child.itertext()), child.get("href", ""))
        elif child.tag in {"strong", "b"}:
            run = paragraph.add_run("".join(child.itertext()))
            run.bold = True
        elif child.tag in {"em", "i"}:
            run = paragraph.add_run("".join(child.itertext()))
            run.italic = True
        else:
            add_inline(paragraph, child)
        if child.tail:
            paragraph.add_run(child.tail)


def add_paragraph_from_node(doc, node, style=None):
    p = doc.add_paragraph(style=style)
    add_inline(p, node)
    return p


def add_list(doc, node, ordered=False):
    style = "List Number" if ordered else "List Bullet"
    for li in node.xpath("./li"):
        p = add_paragraph_from_node(doc, li, style=style)
        p.paragraph_format.left_indent = Inches(0.5)
        p.paragraph_format.first_line_indent = Inches(-0.25)
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.line_spacing = 1.15


def add_table(doc, node):
    rows = node.xpath("./thead/tr | ./tbody/tr | ./tr")
    if not rows:
        return
    columns = max(len(row.xpath("./th|./td")) for row in rows)
    table = doc.add_table(rows=len(rows), cols=columns)
    set_table_borders(table)
    widths = [9360 // columns] * columns
    widths[-1] += 9360 - sum(widths)
    set_table_width(table, widths)
    for r_idx, row in enumerate(rows):
        cells = row.xpath("./th|./td")
        for c_idx, source_cell in enumerate(cells):
            target = table.cell(r_idx, c_idx)
            target.text = ""
            p = target.paragraphs[0]
            add_inline(p, source_cell)
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.line_spacing = 1.15
            if source_cell.tag == "th" or r_idx == 0:
                for run in p.runs:
                    run.bold = True
                p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    doc.add_paragraph().paragraph_format.space_after = Pt(0)


def configure_styles(doc):
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.right_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    normal = doc.styles["Normal"]
    normal.font.name = "Arial"
    set_east_asia_font(normal.font)
    normal.font.size = Pt(11)
    normal.font.color.rgb = RGBColor(0, 0, 0)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(8)
    normal.paragraph_format.line_spacing = 1.15

    for name, size, before, after, color in (
        ("Heading 1", 20, 20, 6, "000000"),
        ("Heading 2", 16, 18, 6, "000000"),
        ("Heading 3", 14, 16, 4, "434343"),
    ):
        style = doc.styles[name]
        style.font.name = "Arial"
        set_east_asia_font(style.font)
        style.font.size = Pt(size)
        style.font.bold = False
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True


def main():
    tree = html.fromstring(SOURCE.read_text(encoding="utf-8"))
    root = tree.xpath('//div[contains(concat(" ", normalize-space(@class), " "), " custom-content ")]')[0]
    doc = Document()
    configure_styles(doc)

    title = doc.add_paragraph()
    title.paragraph_format.space_before = Pt(0)
    title.paragraph_format.space_after = Pt(3)
    run = title.add_run(TITLE)
    run.font.name = "Arial"
    set_east_asia_font(run.font)
    run.font.size = Pt(26)
    run.font.bold = False
    run.font.color.rgb = RGBColor(0, 0, 0)

    meta = doc.add_paragraph()
    meta.paragraph_format.space_after = Pt(12)
    meta_run = meta.add_run(f"投稿先：Marketing\nhandle：ec-site-project-failure-causes\nメタディスクリプション：{META}")
    meta_run.font.name = "Arial"
    set_east_asia_font(meta_run.font)
    meta_run.font.size = Pt(10)
    meta_run.font.color.rgb = RGBColor(85, 85, 85)

    for node in root:
        if node.tag in {"style", "script", "hr"}:
            continue
        if node.tag == "div" and "toc" in (node.get("class") or "").split():
            continue
        if node.tag == "div" and "tbl-scroll" in (node.get("class") or "").split():
            table = node.xpath("./table")
            if table:
                add_table(doc, table[0])
            continue
        if node.tag == "p":
            p = add_paragraph_from_node(doc, node)
            classes = (node.get("class") or "").split()
            if "article-updated" in classes or "article-supervisor" in classes:
                for r in p.runs:
                    r.font.size = Pt(10)
                    r.font.color.rgb = RGBColor(85, 85, 85)
        elif node.tag == "h2":
            add_paragraph_from_node(doc, node, style="Heading 1")
        elif node.tag == "h3":
            add_paragraph_from_node(doc, node, style="Heading 2")
        elif node.tag == "ul":
            add_list(doc, node, ordered=False)
        elif node.tag == "ol":
            add_list(doc, node, ordered=True)
        elif node.tag == "table":
            add_table(doc, node)

    doc.core_properties.title = TITLE
    doc.core_properties.subject = "SOLSTAR Marketingブログ記事下書き"
    doc.core_properties.author = "株式会社SOLSTAR"
    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    main()
